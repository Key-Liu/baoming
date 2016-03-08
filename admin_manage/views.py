# -*- coding: utf-8 -*-
from django.shortcuts import render
from registration.models import *
from math import ceil
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.template.context_processors import csrf
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# 面试时间总批次
T_MAX_NUM = 13
# 每批次面试最少人数
EACH_MIN_NUM = 5
# 每批次面试最懂人数
EACH_MAX_NUM = 10

# 首页
def index_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			return render(request,"admin_manage/index.html")
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

# 计算面试时间页面页面
def compute_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			choice_number = len(Choice.objects.all())
			return render(request,"admin_manage/compute.html",{"choice_number":choice_number})
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

# 计算面试时间
def calculate_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			# 计算各部门的面试总人数
			# 部门列表
			department_list = Department.objects.all()
			# 各部门面试总人数字典，Key为各部门的缩写，Value为各部门面试总人数
			dep_num = {}
			# 测试用
			# result = ""
			for dep in department_list:
				#第一志愿人数
				first_choice_num = len(Choice.objects.filter(first_choice=dep))
				#第二志愿人数
				second_choice_num = len(Choice.objects.filter(second_choice=dep))
				total_num = first_choice_num + second_choice_num
				dep.department_number = total_num
				dep.save()
				dep_num[dep.department_shortcut] = total_num
				# 测试用
				# result = result + dep.department_name + ":" + str(total_num) + "\n"

			# 计算各个部门每批次可分配人员数
			# 各部门各批次可分配人员数字典，Key为各部门的缩写，Value为该部门的面试总人数
			dep_free = {}
			# 测试用
			# result = ""
			for dep in dep_num:
				# 测试用
				# result = result + dep + ":"
				free = {}
				if dep_num[dep] < EACH_MIN_NUM * T_MAX_NUM:
					# 每批次面试人数
					each_num = EACH_MIN_NUM
					# 总批次，取值向上取整
					batch_num = ceil(dep_num[dep] / float(each_num))
				else:
					if dep_num[dep] < EACH_MAX_NUM * T_MAX_NUM:
						each_num = ceil(dep_num[dep] / float(T_MAX_NUM))
					else:
						each_num = EACH_MAX_NUM
					batch_num = T_MAX_NUM

				# 初始化每个批次可分配人数，以批次ID为Key
				i = 0
				while i < batch_num:
					free[i] = each_num
					# 测试用
					# result = result + str(each_num) + " "
					i = i + 1
				# 测试用
				# result = result + "\n"

				# 将各部门各批次可分配人数初始值写入
				dep_free[dep] = free


			# 初次分配
			# 此次分配不考虑时间重叠问题
			for dep in dep_free:
				department = Department.objects.get(department_shortcut=dep)
				# 第一志愿资源
				first_choice_list = Choice.objects.filter(first_choice=department)
				# 标记正在安排的批次
				flag = 0
				for choice in first_choice_list:
					# 当前批次没有容量时
					if dep_free[dep][flag] <= 0:
						flag = flag + 1
					student = choice.student
					try:
						# 检测是否已经安排该学生
						student_interview = Student_interview.objects.get(student=student)
						student_interview.first_dept = department
						student_interview.first_batch = flag
						student_interview.save()
					except Student_interview.DoesNotExist:
						student_interview = Student_interview(student=student,first_dept=department,first_batch=flag)
						student_interview.save()
					dep_free[dep][flag] = dep_free[dep][flag] - 1

				# 第二志愿资源
				second_choice_list = Choice.objects.filter(second_choice=department)
				for choice in second_choice_list:
					if dep_free[dep][flag] <= 0:
						flag = flag + 1
					student = choice.student
					try:
						student_interview = Student_interview.objects.get(student=student)
						student_interview.second_dept = department
						student_interview.second_batch = flag
						student_interview.save()
					except Student_interview.DoesNotExist:
						student_interview = Student_interview(student=student,second_dept=department,second_batch=flag)
						student_interview.save()
					dep_free[dep][flag] = dep_free[dep][flag] - 1

		#######################################
		# 以上代码已测试 #
			
			# 取出时间重叠的人员
			student_interview_list = Student_interview.objects.all()
			for student_interview in student_interview_list:
				# 第一批次时间
				first_batch = student_interview.first_batch
				# 第二批次时间
				second_batch = student_interview.second_batch
				# 时间重叠时
				if first_batch == second_batch:
					student = student_interview.student
					first_dept = student_interview.first_dept
					second_dept = student_interview.second_dept
					other_student_interview = Other_interview(student=student,first_dept=first_dept,second_dept=second_dept,first_batch=first_batch,second_batch=second_batch)
					other_student_interview.save()

			# 该字典存储着各个部门每个批次的学生
			dep_show = {}
			# 获取已经分配学生列表
			student_interview_list = Student_interview.objects.all()
			for student_interview in student_interview_list:
				student = student_interview.student
				# 第一志愿部门
				first_dept = student_interview.first_dept
				# 第一志愿部门批次
				first_batch = student_interview.first_batch
				# 第二志愿部门
				second_dept = student_interview.second_dept
				# 第二志愿部门批次
				second_batch = student_interview.second_batch
				# 当该字典不存在该部门时
				if not dep_show.has_key(first_dept.department_name):
					dep_show[first_dept.department_name] = {}
				# 当该字典不存在该批次时
				if not dep_show[first_dept.department_name].has_key(first_batch):
					dep_show[first_dept.department_name][first_batch] = []
				dep_show[first_dept.department_name][first_batch].append(student)

				if second_dept is not None:
					# 当该字典不存在该部门时
					if not dep_show.has_key(second_dept.department_name):
						dep_show[second_dept.department_name] = {}
					# 当该字典不存在该批次时
					if not dep_show[second_dept.department_name].has_key(second_batch):
						dep_show[second_dept.department_name][second_batch] = []
					dep_show[second_dept.department_name][second_batch].append(student)


			# 操作时间冲突学生
			other_interview_list = Other_interview.objects.all()
			for other_interview in other_interview_list:
				student = other_interview.student
				first_dept = other_interview.first_dept
				second_dept = other_interview.second_dept
				first_batch = other_interview.first_batch
				second_batch = other_interview.second_batch
				# 或缺该学生第一次分配的记录
				student_interview = Student_interview.objects.get(student=student)

				# 两个部门批次均小于2，那么为第二志愿部门添加一个新的批次
				if len(dep_free[first_dept.department_shortcut]) < 2:
					if len(dep_free[second_dept.department_shortcut]) < 2:
						# 计算新批次
						i = 0
						while i in dep_free[second_dept.department_shortcut]:
							i = i + 1
						# 添加新的批次
						dep_free[second_dept.department_shortcut][i] = EACH_MIN_NUM
						# 把该学生安排到新的批次
						dep_show[second_dept.department_name][i] = []
						dep_show[second_dept.department_name][i].append(student)
						# 从原有批次中删除该学生
						dep_show[second_dept.department_name][second_batch].remove(student)
						dep_free[second_dept.department_shortcut][second_batch] = dep_free[second_dept.department_shortcut][second_batch] + 1
						dep_free[second_dept.department_shortcut][i] = dep_free[second_dept.department_shortcut][i] - 1
						student_interview.second_batch = i
						student_interview.save()
						other_interview.delete()
					# 修改第二批志愿批次
					else:
						insert_flag = False
						# 首先判断是否能插入
						for flag in dep_free[second_dept.department_shortcut]:
							# 当前批次不是最后一批
							if (flag + 1) in dep_free[second_dept.department_shortcut]:
								# 当前批次有余量时
								if dep_free[second_dept.department_shortcut][flag] > 0:
									if flag != second_batch:
										# 把该学生安排到新的批次
										dep_show[second_dept.department_name][flag].append(student)
										# 从原有批次中删除该学生
										dep_show[second_dept.department_name][second_batch].remove(student)
										dep_free[second_dept.department_shortcut][flag] = dep_free[second_dept.department_shortcut][flag] - 1
										dep_free[second_dept.department_shortcut][second_batch] = dep_free[second_dept.department_shortcut][second_batch] + 1
										student_interview.second_batch = flag
										student_interview.save()
										other_interview.delete()
										insert_flag = True
										break
							i = i + 1

						if insert_flag == False:
							# 循环以找出除目前被分配批次外的其他最小批次，
							# 并将其设置为该批次进行面试
						
							# 计算最后一批的批次
							i = 0
							while i in dep_show[second_dept.department_name]:
								i = i + 1
							i = i - 1
							# 最后一批的人数
							num = len(dep_show[second_dept.department_name][i])
							# 交换标志，默认为false
							swap_flag = False
							# 从最后一批最后一名学生开始，测试是否能和该学生交换位置
							j = num - 1
							while j >= 0:
								tail_student = dep_show[second_dept.department_name][i][j]
								tail_student_interview = Student_interview.objects.get(student=tail_student)
								if tail_student_interview.first_batch == i:
									if tail_student_interview.second_batch:
										if tail_student_interview.second_batch != second_batch:
											tail_student_interview.first_batch = second_batch
											tail_student_interview.save()
											student_interview.second_batch = i
											student_interview.save()
											dep_show[second_dept.department_name][i].remove(tail_student)
											dep_show[second_dept.department_name][i].append(student)
											dep_show[second_dept.department_name][second_batch].remove(student)
											dep_show[second_dept.department_name][second_batch].append(tail_student)
											swap_flag = True
											other_interview.delete()
											break
									else:
										tail_student_interview.first_batch = second_batch
										tail_student_interview.save()
										student_interview.second_batch = i
										student_interview.save()
										dep_show[second_dept.department_name][i].remove(tail_student)
										dep_show[second_dept.department_name][i].append(student)
										dep_show[second_dept.department_name][second_batch].remove(student)
										dep_show[second_dept.department_name][second_batch].append(tail_student)
										swap_flag = True
										other_interview.delete()
										break
								if tail_student_interview.second_batch == i:
									if tail_student_interview.first_batch != second_batch:
										tail_student_interview.second_batch = second_batch
										tail_student_interview.save()
										student_interview.second_batch = i
										student_interview.save()
										dep_show[second_dept.department_name][i].remove(tail_student)
										dep_show[second_dept.department_name][i].append(student)
										dep_show[second_dept.department_name][second_batch].remove(student)
										dep_show[second_dept.department_name][second_batch].append(tail_student)
										swap_flag = True
										other_interview.delete()
										break	
								j = j - 1

							if swag_flag == False:
								i = i + 1
								# 添加新的批次
								dep_free[second_dept.department_shortcut][i] = EACH_MIN_NUM
								# 把该学生安排到新的批次
								dep_show[second_dept.department_name][i] = []
								dep_show[second_dept.department_name][i].append(student)
								# 从原有批次中删除该学生
								dep_show[second_dept.department_name][second_batch].remove(student)
								dep_free[second_dept.department_shortcut][i] = dep_free[second_dept.department_shortcut][i] - 1
								dep_free[second_dept.department_shortcut][second_batch] = dep_free[second_dept.department_shortcut][second_batch] + 1
								student_interview.second_batch = i
								student_interview.save()
								other_interview.delete()
				# 修改第一批志愿批次
				else:
					insert_flag = False
					# 首先判断是否能插入
					for flag in dep_free[first_dept.department_shortcut]:
						# 当前批次不是最后一批
						if (flag + 1) in dep_free[first_dept.department_shortcut]:
							# 当前批次有余量时
							if dep_free[first_dept.department_shortcut][flag] > 0:
								if flag != first_batch:
									# 把该学生安排到新的批次
									dep_show[first_dept.department_name][flag].append(student)
									# 从原有批次中删除该学生
									dep_show[first_dept.department_name][first_batch].remove(student)
									dep_free[first_dept.department_shortcut][flag] = dep_free[first_dept.department_shortcut][flag] - 1
									dep_free[first_dept.department_shortcut][first_batch] = dep_free[first_dept.department_shortcut][first_batch] + 1
									student_interview.second_batch = flag
									student_interview.save()
									other_interview.delete()
									insert_flag = True
									break

					if insert_flag == False:
						i = 0
						# 计算最后一批的批次
						while i in dep_show[first_dept.department_name]:
							i = i + 1
						i = i - 1
						# 最后一批的人数
						num = len(dep_show[first_dept.department_name][i])
						# 交换标志，默认为false
						swag_flag = False
						# 从最后一批最后一名学生开始，测试是否能和该学生交换位置
						j = num - 1
						while j >= 0:
							tail_student = dep_show[first_dept.department_name][i][j]
							tail_student_interview = Student_interview.objects.get(student=tail_student)
							if tail_student_interview.first_batch == i:
								if tail_student_interview.second_batch:
									if tail_student_interview.second_batch != first_batch:
										tail_student_interview.first_batch = first_batch
										tail_student_interview.save()
										student_interview.first_batch = i
										student_interview.save()
										dep_show[first_dept.department_name][i].remove(tail_student)
										dep_show[first_dept.department_name][i].append(student)
										dep_show[first_dept.department_name][first_batch].remove(student)
										dep_show[first_dept.department_name][first_batch].append(tail_student)
										swag_flag = True
										other_interview.delete()
										break
								else:
									tail_student_interview.first_batch = first_batch
									tail_student_interview.save()
									student_interview.first_batch = i
									student_interview.save()
									dep_show[first_dept.department_name][i].remove(tail_student)
									dep_show[first_dept.department_name][i].append(student)
									dep_show[first_dept.department_name][first_batch].remove(student)
									dep_show[first_dept.department_name][first_batch].append(tail_student)
									swag_flag = True
									other_interview.delete()
									break
							if tail_student_interview.second_batch == i:
								if tail_student_interview.first_batch != first_batch:
									tail_student_interview.second_batch = first_batch
									tail_student_interview.save()
									student_interview.first_batch = i
									student_interview.save()
									dep_show[first_dept.department_name][i].remove(tail_student)
									dep_show[first_dept.department_name][i].append(student)
									dep_show[first_dept.department_name][first_batch].remove(student)
									dep_show[first_dept.department_name][first_batch].append(tail_student)
									swag_flag = True
									other_interview.delete()
									break
							j = j - 1

						if swag_flag == False:
							i = i + 1
							# 添加新的批次
							dep_free[first_dept.department_shortcut][i] = EACH_MIN_NUM
							# 把该学生安排到新的批次
							dep_show[first_dept.department_name][i] = []
							dep_show[first_dept.department_name][i].append(student)
							# 从原有批次中删除该学生
							dep_show[first_dept.department_name][first_batch].remove(student)
							dep_free[first_dept.department_shortcut][first_batch] = dep_free[first_dept.department_shortcut][first_batch] + 1
							dep_free[first_dept.department_shortcut][i] = dep_free[first_dept.department_shortcut][i] - 1
							student_interview.first_batch = i
							student_interview.save()
							other_interview.delete()

			result = "计算成功！"

			return HttpResponse(result)
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

# 展示各部门面试时间页面
def show_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			# 该字典存储着各个部门每个批次的学生
			dep_show = {}
			# 获取已经分配学生列表
			student_interview_list = Student_interview.objects.all()
			for student_interview in student_interview_list:
				student = student_interview.student
				# 第一志愿部门
				first_dept = student_interview.first_dept
				# 第一志愿部门批次
				first_batch = student_interview.first_batch
				# 第二志愿部门
				second_dept = student_interview.second_dept
				# 第二志愿部门批次
				second_batch = student_interview.second_batch
				# 当该字典不存在该部门时
				if not dep_show.has_key(first_dept.department_name):
					dep_show[first_dept.department_name] = {}
				# 当该字典不存在该批次时
				if not dep_show[first_dept.department_name].has_key(first_batch):
					dep_show[first_dept.department_name][first_batch] = []
				dep_show[first_dept.department_name][first_batch].append(student)

				if second_dept is not None:
					# 当该字典不存在该部门时
					if not dep_show.has_key(second_dept.department_name):
						dep_show[second_dept.department_name] = {}
					# 当该字典不存在该批次时
					if not dep_show[second_dept.department_name].has_key(second_batch):
						dep_show[second_dept.department_name][second_batch] = []
					dep_show[second_dept.department_name][second_batch].append(student)

			# 计算所有部门中的最大批次
			batch_max = 0
			for dep in dep_show:
				# 计算该部门总批次
				i = 0
				while i in dep_show[dep]:
					i = i + 1
				i = i - 1
				if i > batch_max:
					batch_max = i
			# 生成批次列表
			dep_batch_list = []
			i = 0
			while i <= batch_max:
				dep_batch_list.append(i)
				i = i + 1

			context = {
				'dep_show' : dep_show,
				'dep_batch_list' : dep_batch_list
			}
			return render(request,'admin_manage/show.html',context)
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

# 计算时间是否重叠页面
def overlap_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			return render(request,"admin_manage/check_overlap.html")
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

# 计算时间是否重叠
def check_overlap_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			# 取出时间重叠的人员
			student_interview_list = Student_interview.objects.all()
			for student_interview in student_interview_list:
				# 第一批次时间
				first_batch = student_interview.first_batch
				# 第二批次时间
				second_batch = student_interview.second_batch
				# 时间重叠时
				if first_batch == second_batch:
					student = student_interview.student
					first_dept = student_interview.first_dept
					second_dept = student_interview.second_dept
					other_student_interview = Other_interview(student=student,first_dept=first_dept,second_dept=second_dept,first_batch=first_batch,second_batch=second_batch)
					other_student_interview.save()
			other_interview_list = Other_interview.objects.all()
			if(len(other_interview_list) == 0):
				return HttpResponse("面试时间正常！")
			else:
				return HttpResponse("面试时间有重叠部分！")
		else:
			return HttpResponse("Not valid account!")
	else:
		return HttpResponse("您当前未登入")

# 返回各个部门的人数
def count_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			# 获取部门列表
			department_list = Department.objects.all()
			dep_num = {}
			for dep in department_list:
				dep_num[dep.department_name] = dep.department_number
			context = {
				'dep_num' : dep_num,
			}
			return render(request,"admin_manage/count.html",context)
		else:
			return HttpResponse("Not valid account!")
	else:
		return HttpResponse("您当前未登入")

# 修改面试时间页面
def check_time_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			c = {}
			c.update(csrf(request))
			return render(request,"admin_manage/change_time.html",c)
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

# 根据学号获取Student_interview对象并返回json数据
@csrf_exempt
def get_stu_inter_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			try:
				# 获取POST过来的Stu_id
				stu_id = request.POST['stu_id']
				stu_object = Student.objects.get(student_id=stu_id)
				stu_interview = Student_interview.objects.get(student=stu_object)
				first_dept_shortcut = stu_interview.first_dept.department_shortcut
				second_dept_shortcut = stu_interview.second_dept.department_shortcut
				first_batch = stu_interview.first_batch
				second_batch = stu_interview.second_batch
				stu_name = stu_object.student_name
				class_num = stu_object.class_num
				doom_num = stu_object.doom_num
				# 返回值
				return_value ={
				"stu_name":stu_name,
				"class_num":class_num,
				"doom_num":doom_num,
				"first_dept":first_dept_shortcut,
				"second_dept":second_dept_shortcut,
				"first_batch":first_batch,
				"second_batch":second_batch
				}
				return HttpResponse(json.dumps(return_value))

			except Exception,ex:
				return HttpResponse()
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")
# 修改面试时间
@csrf_exempt
def save_stu_inter_view(request):
	# 如果已经登入
	if request.user.is_authenticated():
		if request.user.is_staff:
			try:
				# 获取POST过来的Stu_id
				stu_id = request.POST['stu_id']
				first_choice_time = request.POST['first_choice_time']
				second_choice_time = request.POST['second_choice_time']
				stu_object = Student.objects.get(student_id=stu_id)
				stu_interview = Student_interview.objects.get(student=stu_object)
				stu_interview.first_batch = first_choice_time
				stu_interview.second_batch = second_choice_time
				stu_interview.save()
				return HttpResponse("保存成功！")
			except Exception,ex:
				return HttpResponse("保存失败！")
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")

def export_page_view(request):
	return render(request,"admin_manage/export.html")

# 导出操作
def export_doc_view(request):
	if request.user.is_authenticated():
		if request.user.is_staff:
			time_list = []
			time_list.append('9点'.decode("utf-8"))
			time_list.append('9点30分'.decode("utf-8"))
			time_list.append('10点'.decode("utf-8"))
			time_list.append('10点30分'.decode("utf-8"))
			time_list.append('11点'.decode("utf-8"))
			time_list.append('11点30分'.decode("utf-8"))
			time_list.append('2点'.decode("utf-8"))
			time_list.append('2点30分'.decode("utf-8"))
			time_list.append('3点'.decode("utf-8"))
			time_list.append('3点30分'.decode("utf-8"))
			time_list.append('4点'.decode("utf-8"))
			time_list.append('4点30分'.decode("utf-8"))
			time_list.append('5点'.decode("utf-8"))
			student_interview_list= Student_interview.objects.all()
			try:
				for student_interview in student_interview_list:
					# 打开一个空得Document
					document = Document()
					# 添加一个空行
					document.add_paragraph('')

					student_char = '同学：'.decode("utf-8")
					body_char1 = '非常欢迎你报名参加2015年软件学院团委学生会'.decode("utf-8")
					body_char2 = '干事选拔，你被安排于'.decode("utf-8")
					day_char = '日'.decode("utf-8")
					place_char = '地点'.decode("utf-8")
					body_char3 = '进行面试，请务必提前10分钟到达签到，如在时间上安排有任何问题请联系'.decode("utf-8")
					shixiong_char = '师兄/师姐'.decode("utf-8")
					body_char4 = '，谢谢！'.decode("utf-8")
					foot_time_char = '2015年09月11日'.decode("utf-8")
					foot_char = '软件团委学生会'.decode("utf-8")

					paragraph = document.add_paragraph()
					run = paragraph.add_run(student_interview.student.student_name + student_char)
					# 粗体
					run.bold = True
					font = run.font
					font.name = 'SimHei'
					font.size = Pt(21)

					document.add_paragraph('')

					paragraph = document.add_paragraph()
					paragraph_format = paragraph.paragraph_format
					paragraph_format.line_spacing = 1.5
					run = paragraph.add_run('    ' + body_char1 + student_interview.first_dept.department_name + '13' + day_char + time_list[student_interview.first_batch] + 'XXX' + place_char + body_char3 + 'XXX' + shixiong_char + 'XXXXXXXX' + body_char4)
					font = run.font
					font.name = 'SimHei'
					font.size = Pt(14)

					document.add_paragraph('')

					paragraph = document.add_paragraph()
					paragraph_format = paragraph.paragraph_format
					paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
					paragraph_format.line_spacing = 1.5
					run = paragraph.add_run(foot_time_char)
					font = run.font
					font.name = 'SimHei'
					font.size = Pt(14)

					paragraph = document.add_paragraph()
					paragraph_format = paragraph.paragraph_format
					paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
					paragraph_format.line_spacing = 1.5
					run = paragraph.add_run(foot_char)
					font = run.font
					font.name = 'SimHei'
					font.size = Pt(14)

					document.add_paragraph('')
					document.add_paragraph('')
					document.add_paragraph('')

					if student_interview.second_dept:
						paragraph = document.add_paragraph()
						run = paragraph.add_run(student_interview.student.student_name + student_char)
						# 粗体
						run.bold = True
						font = run.font
						font.name = 'SimHei'
						font.size = Pt(21)

						document.add_paragraph('')

						paragraph = document.add_paragraph()
						paragraph_format = paragraph.paragraph_format
						paragraph_format.line_spacing = 1.5
						run = paragraph.add_run('    ' + body_char1 + student_interview.second_dept.department_name + '13' + day_char + time_list[student_interview.second_batch] + 'XXX' + place_char + body_char3 + 'XXX' + shixiong_char + 'XXXXXXXX' + body_char4)
						font = run.font
						font.name = 'SimHei'
						font.size = Pt(14)

						document.add_paragraph('')

						paragraph = document.add_paragraph()
						paragraph_format = paragraph.paragraph_format
						paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
						paragraph_format.line_spacing = 1.5
						run = paragraph.add_run(foot_time_char)
						font = run.font
						font.name = 'SimHei'
						font.size = Pt(14)

						paragraph = document.add_paragraph()
						paragraph_format = paragraph.paragraph_format
						paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
						paragraph_format.line_spacing = 1.5
						run = paragraph.add_run(foot_char)
						font = run.font
						font.name = 'SimHei'
						font.size = Pt(14)

					document.save('doc/' + student_interview.student.student_id + student_interview.student.student_name + '.docx')
				return HttpResponse("导出成功！")
			except Exception,ex:
				return HttpResponse("导出失败！")
		else:
			return HttpResponse("Not valid account!")
	# 用户未登入
	else:
		return HttpResponse("您当前未登入")